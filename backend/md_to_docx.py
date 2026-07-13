"""Convert docs/presentation-strategie.md to Word (.docx)"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

with open('docs/presentation-strategie.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_table = False
table_rows = []

def flush_table(doc, table_rows):
    if not table_rows:
        return
    # Parse header and data
    headers = [c.strip() for c in table_rows[0].split('|')[1:-1]]
    data = []
    for row in table_rows[2:]:
        data.append([c.strip() for c in row.split('|')[1:-1]])
    cols = len(headers)
    table = doc.add_table(rows=1 + len(data), cols=cols)
    table.style = 'Light Grid Accent 1'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(data):
        for j, val in enumerate(row[:cols]):
            table.rows[ri + 1].cells[j].text = val

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Table detection
    if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
        table_rows = [line]
        i += 1
        while i < len(lines) and '|' in lines[i]:
            table_rows.append(lines[i].rstrip('\n'))
            i += 1
        flush_table(doc, table_rows)
        doc.add_paragraph('')
        continue

    # Headings
    if line.startswith('# '):
        h = doc.add_heading(line[2:], level=0)
        i += 1; continue
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
        i += 1; continue
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
        i += 1; continue
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=3)
        i += 1; continue

    # Horizontal rule
    if line.strip() == '---':
        doc.add_paragraph('─' * 50)
        i += 1; continue

    # Code block
    if line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        i += 1  # skip closing ```
        p = doc.add_paragraph()
        p.style = 'No Spacing'
        for cl in code_lines:
            run = p.add_run(cl + '\n')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        doc.add_paragraph('')
        continue

    # Blockquote
    if line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        p.style = 'Intense Quote'
        i += 1; continue

    # Bullet list
    if line.startswith('- '):
        text = line[2:]
        # Bold within bullet
        p = doc.add_paragraph(style='List Bullet')
        # Handle **bold** segments
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1; continue

    # Numbered list
    if re.match(r'^\d+\.\s', line):
        text = re.sub(r'^\d+\.\s', '', line)
        p = doc.add_paragraph(style='List Number')
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1; continue

    # Normal paragraph (handle **bold** and `code`)
    if line.strip():
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
            else:
                p.add_run(part)
        i += 1; continue

    # Empty line
    i += 1

out = 'docs/presentation-strategie.docx'
doc.save(out)
print(f'Saved: {out}')
