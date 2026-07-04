import io
import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


LOGO_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "guardx_logo.png")
NAVY = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x99, 0x99, 0x99)

FOOTER_TEXT = "10 600, boul. Parkway, Montréal (Québec)    H1J 1R6    www.guard-x.com"

MONTHS_FR = [
    "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre",
]

SERVICES = [
    "✓ Inspection et certification de vos extincteurs",
    "✓ Tests fonctionnels de l'éclairage d'urgence (blocs autonomes)",
    "✓ Vérification des systèmes d'alarme incendie et détecteurs",
    "✓ Inspection des systèmes de gicleurs",
]


# ── Low-level helpers ──

def _add_frame(paragraph, x, y, w, h, h_rule="exact"):
    """Inject a w:framePr element for absolute page positioning.

    MUST be called before any other pPr modifications (spacing, indent, etc.)
    so that framePr appears in the correct schema order.
    """
    fp = OxmlElement('w:framePr')
    fp.set(qn('w:x'), str(x))
    fp.set(qn('w:y'), str(y))
    fp.set(qn('w:w'), str(w))
    fp.set(qn('w:h'), str(h))
    fp.set(qn('w:hRule'), h_rule)
    fp.set(qn('w:wrap'), 'notBeside')
    fp.set(qn('w:vAnchor'), 'page')
    fp.set(qn('w:hAnchor'), 'page')
    paragraph._p.get_or_add_pPr().append(fp)


def _set_spacing(paragraph, before=0, after=0, line=None):
    if paragraph is None:
        return
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _set_font(run, size=10.5, bold=False, color=None, name="Arial"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bottom_border(paragraph, color="1F4E79", size="12"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_underline(paragraph):
    for run in paragraph.runs:
        run.font.underline = True


# ── Frame constants (twips, 1440 = 1 inch) ──

# Logo
F_LOGO = dict(x=1095, y=300, w=3168, h=2100, h_rule="atLeast")
# Recipient address — CRITICAL window position
F_ADDR = dict(x=1095, y=2987, w=4410, h=1016, h_rule="exact")
# Objet line
F_OBJET = dict(x=1095, y=4602, w=7650, h=345, h_rule="exact")
# V/Réf line
F_VREF = dict(x=1095, y=5117, w=5655, h=345, h_rule="exact")
# Body text
F_BODY = dict(x=1095, y=5700, w=9735, h=8300, h_rule="atLeast")
# Date — top of page, between logo and address block
F_DATE = dict(x=1095, y=2500, w=9555, h=400, h_rule="exact")
# Footer
F_FOOTER = dict(x=1095, y=15020, w=9735, h=240, h_rule="exact")


# ── POSTAL MODE: absolutely positioned frames for window envelopes ──

def _generate_postal(doc, prospect, settings):
    gestionnaire = (
        prospect.get("Nom_Gestionnaire", "").strip()
        if isinstance(prospect.get("Nom_Gestionnaire"), str) else ""
    )
    syndicat = (
        prospect.get("Nom_Syndicat", "").strip()
        if isinstance(prospect.get("Nom_Syndicat"), str) else ""
    )
    adresse = (
        prospect.get("Adresse", "").strip()
        if isinstance(prospect.get("Adresse"), str) else ""
    )
    ville_cp = (
        prospect.get("Ville_CodePostal", "").strip()
        if isinstance(prospect.get("Ville_CodePostal"), str) else ""
    )

    rep_name = settings.get("rep_name", "")
    rep_title = settings.get("rep_title", "Représentant en protection incendie")
    rep_phone = settings.get("phone", "")
    rep_email = settings.get("email", "")

    # 1 ── LOGO ──
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        _add_frame(p, **F_LOGO)
        _set_spacing(p, 0, 0, 1.0)
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.2))

    # 2 ── RECIPIENT ADDRESS BLOCK (window-critical) ──
    addr_lines = []
    if gestionnaire:
        addr_lines.append(gestionnaire)
    else:
        addr_lines.append("Au président du syndicat de copropriété")
    if syndicat:
        if len(syndicat) > 45:
            syndicat = syndicat[:42] + "..."
        addr_lines.append(syndicat)
    if adresse:
        addr_lines.append(adresse)
    if ville_cp:
        addr_lines.append(ville_cp)

    for line_text in addr_lines[:4]:
        p = doc.add_paragraph(line_text)
        _add_frame(p, **F_ADDR)
        _set_spacing(p, 0, 0, 1.0)
        _set_font(p.runs[0], size=10.5)

    # 3 ── OBJET ──
    p = doc.add_paragraph()
    _add_frame(p, **F_OBJET)
    _set_spacing(p, 0, 0, 1.0)
    run = p.add_run("Objet : Conformité incendie de votre immeuble")
    _set_font(run, size=10.5, bold=True)

    # 4 ── V/RÉF. ──
    ref_text = adresse or ville_cp or ""
    p = doc.add_paragraph()
    _add_frame(p, **F_VREF)
    _set_spacing(p, 0, 0, 1.0)
    run = p.add_run(f"V/Réf. : {ref_text}" if ref_text else "V/Réf. :")
    _set_font(run, size=10)

    # 6 ── BODY TEXT (all paragraphs share identical framePr → same frame) ──

    # Salutation
    if gestionnaire:
        p = doc.add_paragraph(f"Bonjour {gestionnaire},")
    else:
        p = doc.add_paragraph("Bonjour,")
    _add_frame(p, **F_BODY)
    _set_spacing(p, 0, 16, 1.3)
    _set_font(p.runs[0], size=10)

    # Para 1 — gestionnaire / sécurité / RBQ
    p = doc.add_paragraph(
        "À titre de gestionnaire de votre syndicat de copropriété, vous avez la "
        "responsabilité d'assurer la sécurité des résidents. La conformité aux normes "
        "incendie (RBQ et municipales) est essentielle pour protéger votre immeuble "
        "et ses occupants."
    )
    _add_frame(p, **F_BODY)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, 0, 16, 1.3)
    _set_font(p.runs[0], size=10)

    # Para 2 — intro rep / Guard-X / syndicats 8-24
    p = doc.add_paragraph(
        f"Je suis {rep_name}, représentant en protection incendie chez Guard-X, "
        "situé à Anjou. Nous accompagnons les syndicats de copropriété de 8 à 24 "
        "unités dans l'inspection et la certification de leurs équipements de sécurité."
    )
    _add_frame(p, **F_BODY)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, 0, 16, 1.3)
    _set_font(p.runs[0], size=10)

    # Services header
    p = doc.add_paragraph("Nos services incluent :")
    _add_frame(p, **F_BODY)
    _set_spacing(p, 0, 8, 1.3)
    _set_font(p.runs[0], size=10, bold=True)

    # Services list
    for i, svc in enumerate(SERVICES):
        p = doc.add_paragraph(svc)
        _add_frame(p, **F_BODY)
        _set_spacing(p, 0, 6 if i < len(SERVICES) - 1 else 20, 1.3)
        p.paragraph_format.left_indent = Cm(0.5)
        _set_font(p.runs[0], size=10)

    # Engagement + estimation
    p = doc.add_paragraph()
    _add_frame(p, **F_BODY)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, 0, 20, 1.3)
    run = p.add_run("Nous offrons une estimation gratuite et sans engagement. "
                    f"N'hésitez pas à nous contacter au ")
    _set_font(run, size=10)
    run = p.add_run(rep_phone)
    _set_font(run, size=10, bold=True)
    run = p.add_run(" ou par courriel au ")
    _set_font(run, size=10)
    run = p.add_run(rep_email)
    _set_font(run, size=10, bold=True)
    run = p.add_run(".")
    _set_font(run, size=10)

    # Closing
    p = doc.add_paragraph("Merci pour votre temps et votre attention.")
    _add_frame(p, **F_BODY)
    _set_spacing(p, 0, 16, 1.3)
    _set_font(p.runs[0], size=10)

    p = doc.add_paragraph("Cordialement,")
    _add_frame(p, **F_BODY)
    _set_spacing(p, 0, 24, 1.3)
    _set_font(p.runs[0], size=10)

    # Signature
    p = doc.add_paragraph()
    _add_frame(p, **F_BODY)
    _set_spacing(p, 0, 2, 1.15)
    run = p.add_run(rep_name)
    _set_font(run, size=10, bold=True)

    p = doc.add_paragraph(rep_title)
    _add_frame(p, **F_BODY)
    _set_spacing(p, 0, 2, 1.15)
    _set_font(p.runs[0], size=10)

    p = doc.add_paragraph("Votre partenaire de sécurité local à Anjou")
    _add_frame(p, **F_BODY)
    _set_spacing(p, 0, 0, 1.15)
    _set_font(p.runs[0], size=10)

    # 7 ── DATE ──
    today = date.today()
    date_str = f"Le {today.day} {MONTHS_FR[today.month - 1]}, {today.year}"
    p = doc.add_paragraph(date_str)
    _add_frame(p, **F_DATE)
    _set_spacing(p, 0, 0, 1.0)
    _set_font(p.runs[0], size=10.5)

    # 8 ── FOOTER ──
    p = doc.add_paragraph()
    _add_frame(p, **F_FOOTER)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, 0, 0, 1.0)
    run = p.add_run(FOOTER_TEXT)
    _set_font(run, size=8, color=GREY)


# ── DEPOT MODE: flowing layout (no address block, no frames) ──

def _generate_depot(doc, prospect, settings):
    gestionnaire = (
        prospect.get("Nom_Gestionnaire", "").strip()
        if isinstance(prospect.get("Nom_Gestionnaire"), str) else ""
    )
    syndicat = (
        prospect.get("Nom_Syndicat", "").strip()
        if isinstance(prospect.get("Nom_Syndicat"), str) else ""
    )

    rep_name = settings.get("rep_name", "")
    rep_title = settings.get("rep_title", "Représentant en protection incendie")
    rep_phone = settings.get("phone", "")
    rep_email = settings.get("email", "")

    # Logo
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, 0, 2, 1.0)
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.8))

    # Separator
    p = doc.add_paragraph()
    _set_spacing(p, 2, 6, 1.0)
    _add_bottom_border(p, color="1F4E79", size="8")

    # Date
    today = date.today()
    date_str = f"Le {today.day} {MONTHS_FR[today.month - 1]} {today.year}"
    p = doc.add_paragraph(date_str)
    _set_spacing(p, 4, 10, 1.15)

    # Objet
    p = doc.add_paragraph()
    _set_spacing(p, 10, 6, 1.15)
    run = p.add_run("Objet : Conformité incendie de votre immeuble")
    _set_font(run, size=11, bold=True)

    # V/Réf.
    adresse = prospect.get("Adresse", "").strip() if isinstance(prospect.get("Adresse"), str) else ""
    ville_cp = prospect.get("Ville_CodePostal", "").strip() if isinstance(prospect.get("Ville_CodePostal"), str) else ""
    ref_text = adresse or ville_cp or ""
    if ref_text:
        p = doc.add_paragraph()
        _set_spacing(p, 0, 6, 1.15)
        run = p.add_run(f"V/Réf. : {ref_text}")
        _set_font(run, size=10)

    # Salutation
    if gestionnaire:
        p = doc.add_paragraph(f"Bonjour {gestionnaire},")
    else:
        p = doc.add_paragraph("Bonjour,")
    _set_spacing(p, 6, 6, 1.15)

    # Body para 1
    p = doc.add_paragraph(
        "À titre de gestionnaire de votre syndicat de copropriété, vous avez la "
        "responsabilité d'assurer la sécurité des résidents. La conformité aux normes "
        "incendie (RBQ et municipales) est essentielle pour protéger votre immeuble "
        "et ses occupants."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, 0, 6, 1.15)

    # Body para 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, 0, 6, 1.15)
    run = p.add_run(f"Je suis {rep_name}, ")
    _set_font(run, size=11)
    run = p.add_run(rep_title)
    _set_font(run, size=11)
    run = p.add_run(" chez Guard-X, situé à Anjou. Nous accompagnons les syndicats de copropriété de 8 à 24 unités dans l'inspection et la certification de leurs équipements de sécurité.")
    _set_font(run, size=11)

    # Services
    p = doc.add_paragraph()
    _set_spacing(p, 0, 3, 1.15)
    run = p.add_run("Nos services incluent :")
    run.bold = True

    for svc in SERVICES:
        p = doc.add_paragraph(svc)
        _set_spacing(p, 0, 1, 1.15)
        p.paragraph_format.left_indent = Cm(0.5)

    # Engagement
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, 4, 8, 1.15)
    run = p.add_run("Nous offrons une estimation gratuite et sans engagement. N'hésitez pas à nous contacter au ")
    _set_font(run, size=11)
    run = p.add_run(rep_phone)
    _set_font(run, size=11, bold=True)
    run = p.add_run(" ou par courriel au ")
    _set_font(run, size=11)
    run = p.add_run(rep_email)
    _set_font(run, size=11, bold=True)
    run = p.add_run(".")
    _set_font(run, size=11)

    # Closing
    p = doc.add_paragraph("Merci pour votre temps et votre attention.")
    _set_spacing(p, 0, 4, 1.15)

    p = doc.add_paragraph("Cordialement,")
    _set_spacing(p, 0, 20, 1.15)

    # Signature
    p = doc.add_paragraph()
    _set_spacing(p, 0, 2, 1.15)
    run = p.add_run(rep_name)
    _set_font(run, size=11, bold=True)

    p = doc.add_paragraph(rep_title)
    _set_spacing(p, 0, 2, 1.15)
    _set_font(p.runs[0], size=11)

    p = doc.add_paragraph("Votre partenaire de sécurité local à Anjou")
    _set_spacing(p, 0, 0, 1.15)
    _set_font(p.runs[0], size=11)

    # Footer
    p = doc.add_paragraph()
    _set_spacing(p, 12, 0, 1.0)
    _add_horizontal_line(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, 3, 0, 1.0)
    run = p.add_run(FOOTER_TEXT)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ── Main entry point ──

def generate_letter(prospect: dict, settings: dict) -> bytes:
    """Generate a single Word letter and return bytes.

    Postal mode: absolutely positioned frames for window envelopes.
    Depot mode: flowing layout without address block.
    """
    doc = Document()
    mode = settings.get("mode", "postal")

    if mode == "postal":
        # US Letter, all margins 0 — frames handle positioning
        for section in doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)
            section.header_distance = Inches(0)
            section.footer_distance = Inches(0)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)

        _generate_postal(doc, prospect, settings)

    else:
        # Depot mode — flowing layout with normal margins
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(1.8)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        _generate_depot(doc, prospect, settings)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
